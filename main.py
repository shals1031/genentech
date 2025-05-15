"""
Main application module for the Tkinter UI.
This module creates the main application window and connects all UI components.
"""

import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from typing import Dict, Any
import re
from datetime import datetime
import docx
import shutil
import json

# Import custom modules
from data.country_data import COUNTRY_LANGUAGE_DESCRIPTION
from ui import LabeledCombobox, InputField, TextArea
from processor import process_document, process_url, process_image, process_video
# Import the transform function
from ai_service_transform import transform_document_with_openai
from ai_service import format_analysis_document_with_gemini
from processor.document import read_document_file


class Application(tk.Tk):
    """Main application class."""
    analysis_results: str = ""

    def __init__(self):
        """Initialize the application."""
        super().__init__()

        # Configure the main window
        self.title("Pharma Compliance Tool")
        self.geometry("600x500")
        self.minsize(500, 400)

        # Configure the grid layout
        self.columnconfigure(0, weight=1)
        for i in range(5):  # Rows for different components
            self.rowconfigure(i, weight=1 if i == 3 else 0)

        # Initialize UI components
        self._init_ui()

        # Set up event bindings
        self._setup_bindings()

    def _init_ui(self):
        """Initialize all UI components."""
        # Create a frame for the input controls
        self.input_frame = ttk.Frame(self)
        self.input_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=10)
        self.input_frame.columnconfigure(0, weight=1)

        # Country dropdown
        self.country_dropdown = LabeledCombobox(
            self.input_frame,
            "Country:",
            list(COUNTRY_LANGUAGE_DESCRIPTION.keys())
        )
        self.country_dropdown.grid(row=0, column=0, sticky="ew", pady=5)

        # Content type dropdown
        self.content_types = ["URL", "Document", "Image", "Video"]
        self.content_type_dropdown = LabeledCombobox(
            self.input_frame,
            "Content Type:",
            self.content_types,
            on_select=self._on_content_type_changed
        )
        self.content_type_dropdown.grid(row=1, column=0, sticky="ew", pady=5)

        # Input field container (will be populated based on content type)
        self.input_container = ttk.Frame(self)
        self.input_container.grid(row=1, column=0, sticky="ew", padx=10, pady=5)
        self.input_container.columnconfigure(0, weight=1)

        # Initial input field based on default selection (URL)
        self._update_input_field()

        # Process button
        self.process_button = ttk.Button(
            self, 
            text="Process",
            command=self._on_process
        )
        self.process_button.grid(row=2, column=0, pady=10)

        # Output text area
        self.output_area = TextArea(self, "Output:")
        self.output_area.grid(row=3, column=0, sticky="nsew", padx=10, pady=10)

        # Configure tags for colored text
        self.output_area.configure_tag("compliant", foreground="green")
        self.output_area.configure_tag("non_compliant", foreground="red")

    def _setup_bindings(self):
        """Set up event bindings for UI components."""
        # Add any additional bindings here if needed
        pass

    def _on_content_type_changed(self, event):
        """Handle content type selection change."""
        self._update_input_field()

    def _update_input_field(self):
        """Update the input field based on the selected content type."""
        # Clear the current input field
        for widget in self.input_container.winfo_children():
            widget.destroy()

        # Get the selected content type
        content_type = self.content_type_dropdown.get()

        # Create the appropriate input field based on the content type
        if content_type == "URL":
            self.input_field = InputField(
                self.input_container,
                "Enter URL:",
                with_browse_button=False
            )
        elif content_type == "Document":
            # For Document, show the browse button with PDF and TXT file types
            self.input_field = InputField(
                self.input_container,
                f"Select {content_type}:",
                with_browse_button=True,
                filetypes=[("PDF files", "*.pdf"), ("Text files", "*.txt")],
                dialog_title="Select a PDF or TXT file"
            )
        elif content_type == "Image":
            # For Image, show the browse button with image file types
            self.input_field = InputField(
                self.input_container,
                f"Select {content_type}:",
                with_browse_button=True,
                filetypes=[("Image files", "*.jpeg;*.jpg;*.png"), ("JPEG files", "*.jpeg;*.jpg"), ("PNG files", "*.png")],
                dialog_title="Select an image file (JPEG, JPG, PNG)"
            )
        elif content_type == "Video":
            # For Video, show the browse button with video file types
            self.input_field = InputField(
                self.input_container,
                f"Select {content_type}:",
                with_browse_button=True,
                filetypes=[("Video files", "*.mp4;*.webm;*.mkv"), ("MP4 files", "*.mp4"), ("WebM files", "*.webm"), ("MKV files", "*.mkv")],
                dialog_title="Select a video file (MP4, WEBM, MKV)"
            )
        else:
            # For other content types, show the browse button without specific file types
            self.input_field = InputField(
                self.input_container,
                f"Select {content_type}:",
                with_browse_button=True
            )

        self.input_field.grid(row=0, column=0, sticky="ew")

    def _get_ui_values(self):
        """Get values from UI components."""
        country = self.country_dropdown.get()
        language = COUNTRY_LANGUAGE_DESCRIPTION.get(country, "")
        content_type = self.content_type_dropdown.get()
        input_value = self.input_field.get()

        return country, language, content_type, input_value

    def _create_initial_info(self, country, language, content_type, input_value):
        """Create initial processing information text."""
        return (
            f"Processing request:\n"
            f"Country: {country} (Language : {language})\n"
            f"Content Type: {content_type}\n"
            f"Input: {input_value}\n\n"
        )

    def _process_content(self, content_type, input_value, country):
        """Process content based on its type and return the result text."""
        result_text = ""

        if content_type == "URL":
            self.output_area.append("Fetching and analyzing URL content... Please wait.\n")
            for chunk in process_url(
                url=input_value,
                country=country,
            ):
                result_text += chunk
                self.update_idletasks()

        elif content_type == "Document":
            file_type = "PDF" if input_value.lower().endswith('.pdf') else "TXT"
            self.output_area.append(f"Processing {file_type} file... Please wait.\n")
            for chunk in process_document(
                file_path=input_value,
                country=country,
            ):
                result_text += chunk
                self.update_idletasks()

        elif content_type == "Image":
            file_extension = os.path.splitext(input_value)[1].lower()
            file_type = file_extension[1:].upper()
            self.output_area.append(f"Processing {file_type} image... Please wait.\n")
            for chunk in process_image(
                file_path=input_value,
               country=country,
            ):
                result_text += chunk
                self.update_idletasks()

        elif content_type == "Video":
            file_extension = os.path.splitext(input_value)[1].lower()
            file_type = file_extension[1:].upper()
            self.output_area.append(f"Processing {file_type} video... Please wait.\n")
            for chunk in process_video(
                file_path=input_value,
                country=country,
            ):
                result_text += chunk
                self.update_idletasks()

        return result_text

    def _extract_compliance_status(self, result_text, content_type):
        """Extract compliance status from result text."""
        compliance_match = re.search(r'(compliant|non-compliant|not compliant)', result_text, re.IGNORECASE)
        compliance_status = "Compliance status could not be determined"

        if compliance_match:
            status = compliance_match.group(0).lower()
            content_type_text = "content" if content_type == "URL" else content_type.lower()

            if 'not' in status or 'non' in status:
                compliance_status = f"The {content_type_text} is NOT COMPLIANT with medical norms."
            else:
                compliance_status = f"The {content_type_text} is COMPLIANT with medical norms."

        return compliance_status

    def _create_and_save_json_document(self, result_text, country, input_value):
        """Create and save a JSON document with the analysis results."""
        json_iterator = format_analysis_document_with_gemini(result_text, country)

        # Collect all chunks from the iterator to form a complete JSON string
        json_result = ""
        for chunk in json_iterator:
            json_result += chunk

        # Extract JSON from the response if it's wrapped in markdown code blocks
        json_string = self._extract_json_from_text(json_result)

        # Validate and format the JSON
        try:
            # Parse the JSON to validate it
            if not json_string:
                raise json.JSONDecodeError("No valid JSON found in the response", "", 0)
            parsed_json = json.loads(json_string)
            # Format it with indentation for better readability
            json_result = json.dumps(parsed_json, indent=4)

            # Save the JSON to a file on the client side
            base_filename = os.path.splitext(os.path.basename(input_value))[0] if os.path.isfile(input_value) else "analysis"
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"{base_filename}_analysis_{timestamp}.json"

            save_path = filedialog.asksaveasfilename(
                defaultextension=".json",
                filetypes=[("JSON files", "*.json")],
                initialfile=default_filename
            )

            if save_path:
                with open(save_path, 'w', encoding='utf-8') as json_file:
                    json_file.write(json_result)
                self.output_area.append(f"\nAnalysis JSON saved to: {save_path}")
            else:
                self.output_area.append("\nJSON save operation cancelled. Analysis not saved to file.")
        except json.JSONDecodeError as e:
            error_msg = f"Error parsing JSON: {str(e)}"
            messagebox.showerror("JSON Error", error_msg)
            self.output_area.append(f"\n{error_msg}")
            self.output_area.append("\nRaw response will not be saved as JSON.")

    def _create_and_save_document(self, result_text, compliance_status, input_value, country):
        """Create and save a Word document with the analysis results."""
        doc = docx.Document()

        # Create heading based on content type
        if input_value.startswith(('http://', 'https://')):
            doc.add_heading(f'Compliance Analysis for URL: {input_value}', 0)
            # Clean URL for filename
            url_filename = re.sub(r'[^\w]', '_', input_value)[:30]
            base_filename = f"url_{url_filename}"
        else:
            doc.add_heading(f'Compliance Analysis for {os.path.basename(input_value)}', 0)
            base_filename = os.path.splitext(os.path.basename(input_value))[0]

        # Add common document content
        doc.add_paragraph(f'Country: {country}')
        doc.add_paragraph(f'Analysis Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Compliance Status: {compliance_status}')
        doc.add_heading('Detailed Analysis', level=1)
        doc.add_paragraph(result_text)

        # Generate filename and ask the user where to save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_filename = f"{base_filename}_analysis_{timestamp}.docx"

        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            initialfile=default_filename
        )

        saved_message = ""
        if save_path:
            doc.save(save_path)
            saved_message = f"\nDetailed analysis saved to: {save_path}"
        else:
            saved_message = "\nSave operation cancelled. Detailed analysis not saved."

        return saved_message

    def _update_ui_with_results(self, initial_info, compliance_status,  input_value=None):
        """Update the UI with processing results."""
        self.output_area.set(initial_info)

        # Apply appropriate color tag based on compliance status
        # Only show transform button if the document is non compliant
        if "NOT COMPLIANT" in compliance_status:
            self.output_area.append(f"\n{compliance_status}\n", "non_compliant")
            # Store the current document path for transformation
            self.current_document_path = input_value

            # Create a transform button
            self.transform_button = ttk.Button(
                self,
                text="Transform to Compliant Document",
                command=self._on_transform
            )
            self.transform_button.grid(row=4, column=0, pady=10)
        else:
            self.output_area.append(f"\n{compliance_status}\n", "compliant")

            # Remove transform button if it exists
            if hasattr(self, 'transform_button') and self.transform_button:
                self.transform_button.destroy()
                self.transform_button = None

        # self.output_area.append("\nPlease view the downloaded file for detailed analysis.")
        self.output_area.append("\n\nProcessing complete!")

    def _extract_json_from_text(self, text):
        """Extract JSON from text that might contain markdown or other formatting."""
        # Try to find JSON between triple backticks (markdown code blocks)
        import re
        json_pattern = r'```(?:json)?\s*([\s\S]*?)\s*```'
        matches = re.findall(json_pattern, text)

        if matches:
            # Try each match until we find valid JSON
            for match in matches:
                try:
                    # Validate that this is actually JSON
                    json.loads(match.strip())
                    return match.strip()
                except json.JSONDecodeError:
                    continue

        # If no valid JSON found in code blocks, try to find JSON objects directly
        # Look for text that starts with { and ends with }
        json_pattern = r'(\{[\s\S]*\})'
        matches = re.findall(json_pattern, text)

        if matches:
            # Try each match until we find valid JSON
            for match in matches:
                try:
                    # Validate that this is actually JSON
                    json.loads(match.strip())
                    return match.strip()
                except json.JSONDecodeError:
                    continue

        # If we still haven't found valid JSON, return the original text
        # (it will likely fail JSON parsing, but we'll handle that in the calling function)
        return text.strip()

    def _is_valid_input(self, content_type, input_value):
        """Check if the input is valid for the selected content type."""
        if content_type == "URL" and input_value:
            return True
        elif content_type == "Document" and (input_value.lower().endswith('.pdf') or input_value.lower().endswith('.txt')):
            return True
        elif content_type == "Image" and any(input_value.lower().endswith(ext) for ext in ['.jpeg', '.jpg', '.png']):
            return True
        elif content_type == "Video" and any(input_value.lower().endswith(ext) for ext in ['.mp4', '.webm', '.mkv']):
            return True
        return False

    def _on_process(self):
        """Handle process button click."""
        # Get values from UI components
        country, language, content_type, input_value = self._get_ui_values()

        # Create initial info and display it
        initial_info = self._create_initial_info(country, language, content_type, input_value)
        self.output_area.set(initial_info)
        self.update_idletasks()

        # Check if input is valid for the selected content type
        if self._is_valid_input(content_type, input_value):
            try:
                # Process the content based on its type
                self.output_area.append(f"Processing {content_type}...\n")
                self.update_idletasks()

                result_text = self._process_content(content_type, input_value,country)

                self.analysis_text = result_text

                # Extract compliance status
                compliance_status = self._extract_compliance_status(result_text, content_type)

                # Create and save a JSON document
                self._create_and_save_json_document(result_text, country, input_value)

                # Create and save document
                # saved_message = self._create_and_save_document(result_text, compliance_status, input_value, country)

                # Update UI with results
                self._update_ui_with_results(initial_info, compliance_status , input_value)

            except FileNotFoundError:
                messagebox.showerror("Error", f"File not found: {input_value}")
                self.output_area.append("Error: File not found.")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {str(e)}")
                self.output_area.append(f"Error: {str(e)}")
        else:
            # For invalid inputs, display a placeholder message
            self.output_area.append("This functionality is not implemented yet or input is invalid.")
            self.output_area.append("Processing complete!")

    def _on_transform(self):
        """Handle transform button click."""
        if not hasattr(self, 'current_document_path') or not self.current_document_path:
            messagebox.showerror("Error", "No document to transform.")
            return

        try:
            # Get the country from the UI
            country = self.country_dropdown.get()

            # Read the document file
            document_data, file_type = read_document_file(self.current_document_path)

            # Update UI
            self.output_area.append("\nTransforming document to be compliant... Please wait.\n")
            self.update_idletasks()



            # Transform the document
            transformed_pdf_path = transform_document_with_openai(self.analysis_text,document_data, file_type, country)

            # Ask user where to save the transformed PDF
            base_filename = os.path.splitext(os.path.basename(self.current_document_path))[0]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"{base_filename}_transformed_{timestamp}.pdf"

            save_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                initialfile=default_filename
            )

            if save_path:
                # Copy the temporary PDF to the user-selected location
                import shutil
                shutil.copy2(transformed_pdf_path, save_path)

                # Remove the temporary file
                os.remove(transformed_pdf_path)

                # Update UI
                self.output_area.append(f"\nTransformed document saved to: {save_path}")
                self.output_area.append("\nTransformation complete!")
            else:
                # Remove the temporary file
                os.remove(transformed_pdf_path)
                self.output_area.append("\nSave operation cancelled. Transformed document not saved.")

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred during transformation: {str(e)}")
            self.output_area.append(f"Error: {str(e)}")


def main():
    """Main function to start the application."""
    app = Application()
    app.mainloop()


if __name__ == "__main__":
    main()
